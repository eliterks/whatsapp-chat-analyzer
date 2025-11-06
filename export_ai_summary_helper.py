# ==========================================
# export_ai_summary_helper.py
# ==========================================
import io
import json
import pandas as pd
from datetime import datetime
from reportlab.lib.pagesizes import A4
from reportlab.lib import colors
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle
from reportlab.lib.styles import getSampleStyleSheet
from docx import Document

# ------------------------------------------------
# 📄 Export as CSV
# ------------------------------------------------
def export_ai_summary_csv(summary_text, avg_sentiment, mood, top_topics):
    buffer = io.StringIO()
    topics_str = ", ".join([t[0] for t in top_topics])
    buffer.write("Summary,Mood,Average Sentiment,Top Topics,Date\n")
    buffer.write(f'"{summary_text.strip()}","{mood}",{avg_sentiment},"{topics_str}","{datetime.now().strftime("%Y-%m-%d")}"\n')
    return buffer.getvalue()

# ------------------------------------------------
# 📘 Export as Excel
# ------------------------------------------------
def export_ai_summary_excel(summary_text, avg_sentiment, mood, top_topics):
    output = io.BytesIO()
    topics_df = pd.DataFrame(top_topics, columns=["Topic", "Mentions"])
    summary_df = pd.DataFrame({
        "Summary": [summary_text],
        "Mood": [mood],
        "Average Sentiment": [avg_sentiment],
        "Date": [datetime.now().strftime("%Y-%m-%d")]
    })

    with pd.ExcelWriter(output, engine="openpyxl") as writer:
        summary_df.to_excel(writer, index=False, sheet_name="Summary")
        topics_df.to_excel(writer, index=False, sheet_name="Top Topics")

    return output.getvalue()

# ------------------------------------------------
# 📄 Export as PDF
# ------------------------------------------------
def export_ai_summary_pdf(summary_text, avg_sentiment, mood, top_topics):
    buffer = io.BytesIO()
    doc = SimpleDocTemplate(buffer, pagesize=A4)
    styles = getSampleStyleSheet()
    story = []

    story.append(Paragraph("📱 WhatsApp Chat AI Summary Report", styles["Title"]))
    story.append(Spacer(1, 12))

    story.append(Paragraph(f"<b>Date:</b> {datetime.now().strftime('%Y-%m-%d')}", styles["Normal"]))
    story.append(Paragraph(f"<b>Mood:</b> {mood}", styles["Normal"]))
    story.append(Paragraph(f"<b>Average Sentiment:</b> {avg_sentiment:.3f}", styles["Normal"]))
    story.append(Spacer(1, 12))

    story.append(Paragraph("<b>Summary:</b>", styles["Heading2"]))
    story.append(Paragraph(summary_text, styles["Normal"]))
    story.append(Spacer(1, 12))

    story.append(Paragraph("<b>Top Discussion Topics:</b>", styles["Heading2"]))
    data = [["Topic", "Mentions"]] + [[t[0], str(t[1])] for t in top_topics]
    table = Table(data, colWidths=[250, 100])
    table.setStyle(TableStyle([
        ('BACKGROUND', (0,0), (-1,0), colors.lightblue),
        ('TEXTCOLOR', (0,0), (-1,0), colors.black),
        ('ALIGN', (0,0), (-1,-1), 'CENTER'),
        ('FONTNAME', (0,0), (-1,0), 'Helvetica-Bold'),
        ('BOTTOMPADDING', (0,0), (-1,0), 10),
        ('GRID', (0,0), (-1,-1), 0.5, colors.grey),
    ]))
    story.append(table)

    doc.build(story)
    return buffer.getvalue()

# ------------------------------------------------
# 💾 Export as JSON
# ------------------------------------------------
def export_ai_summary_json(summary_text, avg_sentiment, mood, top_topics):
    data = {
        "summary": summary_text.strip(),
        "mood": mood,
        "average_sentiment": avg_sentiment,
        "top_topics": [{"topic": t[0], "mentions": t[1]} for t in top_topics],
        "date": datetime.now().strftime("%Y-%m-%d")
    }
    return json.dumps(data, indent=4)

# ------------------------------------------------
# 📝 Export as Word
# ------------------------------------------------
def export_ai_summary_word(summary_text, avg_sentiment, mood, top_topics):
    buffer = io.BytesIO()
    doc = Document()
    doc.add_heading("WhatsApp Chat AI Summary Report", level=1)

    doc.add_paragraph(f"Date: {datetime.now().strftime('%Y-%m-%d')}")
    doc.add_paragraph(f"Mood: {mood}")
    doc.add_paragraph(f"Average Sentiment: {avg_sentiment:.3f}")
    doc.add_heading("Summary:", level=2)
    doc.add_paragraph(summary_text)

    doc.add_heading("Top Discussion Topics:", level=2)
    for topic, count in top_topics:
        doc.add_paragraph(f"{topic} — {count} mentions")

    doc.save(buffer)
    return buffer.getvalue()

# ------------------------------------------------
# 🌟 Streamlit Helper for Download Buttons
# ------------------------------------------------
def show_download_buttons(summary_text, avg_sentiment, mood, top_topics):
    import streamlit as st
    st.subheader("📤 Export AI Summary Results")

    col1, col2, col3, col4, col5 = st.columns(5)

    with col1:
        csv_data = export_ai_summary_csv(summary_text, avg_sentiment, mood, top_topics)
        st.download_button("⬇️  📄 Download CSV", csv_data, "ai_summary.csv", "text/csv")

    with col2:
        excel_data = export_ai_summary_excel(summary_text, avg_sentiment, mood, top_topics)
        st.download_button("⬇️  📘 Download Excel", excel_data, "ai_summary.xlsx",
                           "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet")

    with col3:
        pdf_data = export_ai_summary_pdf(summary_text, avg_sentiment, mood, top_topics)
        st.download_button("⬇️  🧾 Download PDF", pdf_data, "ai_summary.pdf", "application/pdf")

    with col4:
        json_data = export_ai_summary_json(summary_text, avg_sentiment, mood, top_topics)
        st.download_button("⬇️  💾 Download JSON", json_data, "ai_summary.json", "application/json")

    with col5:
        word_data = export_ai_summary_word(summary_text, avg_sentiment, mood, top_topics)
        st.download_button("⬇️  📝 Download Word", word_data, "ai_summary.docx",
                           "application/vnd.openxmlformats-officedocument.wordprocessingml.document")
