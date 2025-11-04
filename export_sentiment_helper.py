"""
===========================================================
📦 export_sentiment_helper.py — Sentiment Export Utilities
===========================================================

Exports sentiment analysis results to CSV, Excel, PDF, and Word.
Compatible with Streamlit-based WhatsApp Chat Analyzer.

Dependencies:
    pip install reportlab pillow openpyxl python-docx
===========================================================
"""

import io
import pandas as pd
import re
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle
from reportlab.lib import colors
from reportlab.lib.pagesizes import A4
from reportlab.lib.styles import getSampleStyleSheet
from docx import Document
from datetime import datetime
from xml.sax.saxutils import escape as html_escape


# ==============================
# 🧹 SANITIZATION HELPER
# ==============================
def sanitize_for_excel(text):
    """
    Remove illegal characters that OpenPyXL cannot handle in Excel worksheets.
    OpenPyXL doesn't allow control characters (0x00-0x1F except tab, newline, carriage return).
    
    Args:
        text: String to sanitize
        
    Returns:
        Sanitized string safe for Excel
    """
    if not isinstance(text, str):
        return text
    
    # Remove control characters except \t (tab), \n (newline), \r (carriage return)
    # Control characters range: 0x00-0x1F
    # Keep: \t (0x09), \n (0x0A), \r (0x0D)
    sanitized = re.sub(r'[\x00-\x08\x0B-\x0C\x0E-\x1F]', '', text)
    
    # Also remove other problematic characters
    # Remove vertical tab, form feed, and other unusual whitespace
    sanitized = re.sub(r'[\x7F-\x9F]', '', sanitized)
    
    # Remove Unicode directional formatting characters (cause issues in Excel)
    # These include: LRM, RLM, LRE, RLE, PDF, LRO, RLO, etc.
    sanitized = re.sub(r'[\u200E\u200F\u202A-\u202E\u2066-\u2069]', '', sanitized)
    
    # Remove additional problematic Unicode characters that can cause issues
    # Zero-width characters and other invisible characters
    sanitized = re.sub(r'[\u200B-\u200D\uFEFF]', '', sanitized)
    
    # Remove variation selectors and other combining characters that might cause issues
    sanitized = re.sub(r'[\uFE00-\uFE0F]', '', sanitized)
    
    # Additional safety: remove any character that might cause issues with OpenPyXL
    # This is a more aggressive approach to ensure compatibility
    try:
        # Test if the string can be written to Excel
        import io
        import pandas as pd
        test_df = pd.DataFrame({'test': [sanitized]})
        buffer = io.BytesIO()
        with pd.ExcelWriter(buffer, engine='openpyxl') as writer:
            test_df.to_excel(writer, index=False)
        buffer.close()
    except Exception:
        # If it fails, remove all non-ASCII characters as a last resort
        sanitized = re.sub(r'[^\x00-\x7F]', '', sanitized)
    
    return sanitized


# ==============================
# 1️⃣ CSV EXPORT
# ==============================
def export_sentiment_csv(sentiment_df: pd.DataFrame) -> bytes:
    csv_data = sentiment_df.to_csv(index=False).encode("utf-8")
    return csv_data


# ==============================
# 2️⃣ EXCEL EXPORT
# ==============================
def export_sentiment_excel(sentiment_summary: dict, sentiment_df: pd.DataFrame) -> bytes:
    output = io.BytesIO()
    
    # Create a copy of the DataFrame to avoid modifying the original
    df_clean = sentiment_df.copy()
    
    # Sanitize all string-like cells regardless of dtype (covers object and pandas StringDtype)
    df_clean = df_clean.applymap(lambda x: sanitize_for_excel(x) if isinstance(x, str) else x)
    
    with pd.ExcelWriter(output, engine='openpyxl') as writer:
        # Summary Sheet
        summary_df = pd.DataFrame({
            "Metric": ["Positive", "Neutral", "Negative", "Average Sentiment"],
            "Value": [
                sentiment_summary["positive"],
                sentiment_summary["neutral"],
                sentiment_summary["negative"],
                round(sentiment_summary["avg_sentiment"], 3)
            ]
        })
        summary_df.to_excel(writer, index=False, sheet_name="Summary")

        # Messages Sheet (with sanitized data)
        df_clean.to_excel(writer, index=False, sheet_name="Messages")
    
    output.seek(0)
    return output.getvalue()


# ==============================
# 3️⃣ PDF EXPORT
# ==============================
def export_sentiment_pdf(selected_user: str, sentiment_summary: dict, sentiment_df: pd.DataFrame) -> bytes:
    buffer = io.BytesIO()
    doc = SimpleDocTemplate(buffer, pagesize=A4)
    styles = getSampleStyleSheet()
    elements = []

    elements.append(Paragraph(f"<b>WhatsApp Sentiment Analysis Report</b>", styles["Title"]))
    elements.append(Spacer(1, 12))
    elements.append(Paragraph(f"<b>User:</b> {selected_user}", styles["Normal"]))
    elements.append(Paragraph(f"<b>Date:</b> {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}", styles["Normal"]))
    elements.append(Spacer(1, 12))

    data = [
        ["Sentiment", "Count"],
        ["😊 Positive", sentiment_summary["positive"]],
        ["😐 Neutral", sentiment_summary["neutral"]],
        ["😡 Negative", sentiment_summary["negative"]],
        ["📈 Avg Sentiment", round(sentiment_summary["avg_sentiment"], 3)],
    ]
    table = Table(data, colWidths=[150, 150])
    table.setStyle(TableStyle([
        ("BACKGROUND", (0, 0), (-1, 0), colors.lightblue),
        ("TEXTCOLOR", (0, 0), (-1, 0), colors.whitesmoke),
        ("ALIGN", (0, 0), (-1, -1), "CENTER"),
        ("GRID", (0, 0), (-1, -1), 1, colors.grey)
    ]))
    elements.append(table)
    elements.append(Spacer(1, 18))

    elements.append(Paragraph("<b>Sample Messages:</b>", styles["Heading2"]))
    # Escape and sanitize text to avoid ReportLab/HTML issues
    for _, row in sentiment_df.head(10).iterrows():
        try:
            sender = html_escape(str(row.get('Sender', 'Unknown')))
            message = html_escape(sanitize_for_excel(str(row.get('Message', ''))))
            sentiment = html_escape(str(row.get('Sentiment', '')))
            polarity = float(row.get('Polarity', 0.0))
            # Build a safe paragraph string
            msg = f"<b>{sender}</b>: {message} (<i>{sentiment}</i>, Polarity={polarity:.3f})"
        except Exception:
            # Fallback to a simpler safe string if anything goes wrong
            msg = html_escape(str(row))
        elements.append(Paragraph(msg, styles["Normal"]))
        elements.append(Spacer(1, 6))

    # Build the PDF; be defensive against encoding/layout errors
    try:
        doc.build(elements)
    except Exception:
        # Fallback: build a minimal PDF without messages if rendering fails
        elements_fallback = elements[:5]
        elements_fallback.append(Paragraph("(Message rendering truncated due to encoding issues)", styles["Italic"]))
        doc = SimpleDocTemplate(buffer, pagesize=A4)
        doc.build(elements_fallback)
    buffer.seek(0)
    return buffer.getvalue()


# ==============================
# 4️⃣ WORD EXPORT
# ==============================
def export_sentiment_word(selected_user: str, sentiment_summary: dict, sentiment_df: pd.DataFrame) -> bytes:
    doc = Document()
    doc.add_heading("WhatsApp Sentiment Analysis Report", level=1)
    doc.add_paragraph(f"User: {selected_user}")
    doc.add_paragraph(f"Date Generated: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")

    doc.add_heading("Sentiment Summary", level=2)
    table = doc.add_table(rows=1, cols=2)
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Sentiment'
    hdr_cells[1].text = 'Value'

    for key, value in sentiment_summary.items():
        row_cells = table.add_row().cells
        row_cells[0].text = key.capitalize()
        row_cells[1].text = str(round(value, 3) if isinstance(value, float) else value)

    doc.add_heading("Sample Messages", level=2)
    for _, row in sentiment_df.head(10).iterrows():
        doc.add_paragraph(f"{row['Sender']}: {row['Message']} "
                          f"({row['Sentiment']}, Polarity={row['Polarity']:.3f})")

    output = io.BytesIO()
    doc.save(output)
    output.seek(0)
    return output.getvalue()
